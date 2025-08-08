import streamlit as st
import pandas as pd
import altair as alt
import random
import time
import re
from io import BytesIO
from pandas import ExcelWriter
from docx import Document

# -----------------------------
# Helper Functions
# -----------------------------

def mock_parse_cv_and_jd():
    return {
        "educationScore": random.choice([10, 7, 4, 2]),
        "experienceScore": random.choice([10, 7, 5, 3])
    }

def mock_parse_interview_sheet():
    return {
        "performanceScore": random.choice([10, 7, 5, 2])
    }

def get_step_interval(score: int):
    if score >= 25:
        return list(range(12, 16)), "Top Range"
    elif score >= 20:
        return list(range(9, 12)), "Mid-Upper Range"
    elif score >= 15:
        return list(range(6, 9)), "Mid Range"
    elif score >= 10:
        return list(range(3, 6)), "Lower-Mid Range"
    else:
        return list(range(1, 3)), "Bottom Range"

def convert_df_to_excel(df: pd.DataFrame) -> bytes:
    output = BytesIO()
    with ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="EquityAnalysis", index=False)
    return output.getvalue()

def load_filtered_equity_data(uploaded_file, position_title_input: str):
    try:
        df = pd.read_excel(uploaded_file)
    except Exception as e:
        return None, f"❌ Could not read Excel file: {e}"

    df.columns = [col.strip().lower() for col in df.columns]

    required_cols = {'id', 'position title', 'hire date', 'comp rate'}
    if not required_cols.issubset(set(df.columns)):
        return None, "❌ Excel must include columns: ID, Position Title, Hire Date, Comp Rate."

    df_filtered = df[[c for c in df.columns if c in required_cols]].copy()
    df_filtered.columns = ['id', 'positionTitle', 'hireDate', 'compRate']

    if not position_title_input:
        return df_filtered.iloc[0:0], None

    df_filtered = df_filtered[
        df_filtered['positionTitle'].astype(str).str.lower() == position_title_input.lower()
    ]
    return df_filtered, None

# ---------- Salary Scale helpers ----------

def _extract_step_number(col_name: str):
    """
    Accepts headers like 'Minimum 1', '1', 'Medium 8', 'Maximum 15', 'Step 3', etc.
    Returns an int step or None if no match.
    """
    s = str(col_name).strip().lower()
    m = re.findall(r'(\d+)', s)
    if not m:
        return None
    return int(m[-1])

def load_salary_scale(scale_file):
    """
    Supports:
      A) Long table: columns Band, Grade, Step, Amount
      B) Wide table: columns Band, Grade, [Step % optional], and step columns (e.g., 'Minimum 1', '2', ..., 'Maximum 15')
    Returns normalized DataFrame with: band(str), grade(Int), step(Int), amount(float)
    """
    try:
        df = pd.read_excel(scale_file)
    except Exception as e:
        return None, f"❌ Could not read Salary Scale file: {e}"

    df.columns = [str(c).strip() for c in df.columns]

    try:
        band_col = df.columns[[c.lower() == 'band' for c in df.columns]][0]
        grade_col = df.columns[[c.lower() == 'grade' for c in df.columns]][0]
    except IndexError:
        return None, "❌ Salary Scale must include 'Band' and 'Grade' columns."

    if {'step', 'amount'}.issubset(set(c.lower() for c in df.columns)):
        step_col = df.columns[[c.lower() == 'step' for c in df.columns]][0]
        amount_col = df.columns[[c.lower() == 'amount' for c in df.columns]][0]
        out = df[[band_col, grade_col, step_col, amount_col]].copy()
        out.columns = ['band', 'grade', 'step', 'amount']
        out['band'] = out['band'].astype(str).str.strip().str.upper()
        out['grade'] = pd.to_numeric(out['grade'], errors='coerce').astype('Int64')
        out['step'] = pd.to_numeric(out['step'], errors='coerce').astype('Int64')
        out['amount'] = pd.to_numeric(out['amount'], errors='coerce')
        out = out.dropna(subset=['grade', 'step', 'amount'])
        return out, None
    else:
        step_cols = []
        for c in df.columns:
            if c in [band_col, grade_col]:
                continue
            step_num = _extract_step_number(c)
            if step_num is not None:
                step_cols.append((c, step_num))

        if not step_cols:
            return None, "❌ Could not detect step columns in Salary Scale. Provide 'Step/Amount' or columns like '1..15'."

        melt_cols = [band_col, grade_col] + [c for c, _ in step_cols]
        wide = df[melt_cols].copy()
        long_df = wide.melt(id_vars=[band_col, grade_col], var_name='step_col', value_name='amount')
        long_df['step'] = long_df['step_col'].apply(_extract_step_number)
        long_df = long_df.drop(columns=['step_col'])

        long_df.columns = ['band', 'grade', 'amount', 'step']
        long_df['band'] = long_df['band'].astype(str).str.strip().str.upper()
        long_df['grade'] = pd.to_numeric(long_df['grade'], errors='coerce').astype('Int64')
        long_df['step'] = pd.to_numeric(long_df['step'], errors='coerce').astype('Int64')
        long_df['amount'] = pd.to_numeric(long_df['amount'], errors='coerce')
        long_df = long_df.dropna(subset=['grade', 'step', 'amount'])

        return long_df[['band', 'grade', 'step', 'amount']], None

def parse_band_grade(grade_text: str):
    """
    Accepts formats like 'E-6', 'E 6', 'e6', 'Band E Grade 6'
    Returns (band:str, grade:int) or (None, None)
    """
    if not grade_text:
        return None, None
    s = str(grade_text).strip()
    m = re.search(r'([A-Za-z])\s*[-– ]?\s*(\d+)', s)
    if not m:
        return None, None
    band = m.group(1).upper()
    grade = int(m.group(2))
    return band, grade

# ----------------------------------------------

def generate_word_report(name, title, grade, education_score, experience_score, performance_score,
                         total_score, interval_options, placement, selected_step,
                         recommended_salary, final_salary, budget_threshold, hr_comments, df_peers):

    doc = Document()
    doc.add_heading('Final Salary Evaluation Report', 0)

    doc.add_heading('Candidate Details', level=1)
    doc.add_paragraph(f"Name: {name or 'N/A'}")
    doc.add_paragraph(f"Position Title: {title or 'N/A'}")
    doc.add_paragraph(f"Grade: {grade or 'N/A'}")

    doc.add_heading('Scoring Breakdown', level=1)
    doc.add_paragraph(f"Education Score: {education_score}/10")
    doc.add_paragraph(f"Experience Score: {experience_score}/10")
    doc.add_paragraph(f"Performance Score: {performance_score}/10")
    doc.add_paragraph(f"Total Score: {total_score}/30")
    doc.add_paragraph(f"Suggested Step Interval: {interval_options} → {placement}")
    doc.add_paragraph(f"Final Selected Step: {selected_step}")

    doc.add_heading('Salary Recommendation', level=1)
    doc.add_paragraph(f"AI-Recommended Salary: AED {recommended_salary:,.0f}")
    doc.add_paragraph(f"Final Recommended Salary: AED {final_salary:,.0f}")
    doc.add_paragraph(f"Budget Threshold: AED {budget_threshold:,.0f}")
    budget_status = "Within Budget ✅" if final_salary <= budget_threshold else "Out of Budget ❌"
    doc.add_paragraph(f"Budget Status: {budget_status}")

    if isinstance(df_peers, pd.DataFrame) and not df_peers.empty:
        doc.add_heading('Internal Equity Data', level=1)
        table = doc.add_table(rows=1, cols=len(df_peers.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df_peers.columns):
            hdr_cells[i].text = str(col_name)

        for _, row in df_peers.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

    doc.add_heading("HR Final Comments", level=1)
    doc.add_paragraph(hr_comments.strip() if hr_comments and hr_comments.strip() else "N/A")
    return doc

# -----------------------------
# Streamlit UI
# -----------------------------

st.set_page_config(page_title="Salary Evaluation Dashboard", layout="wide")
st.markdown("<h1 style='color:#003366;'>📊 Salary Evaluation Dashboard</h1>", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["📘 Evaluation & Scoring Matrices", "📋 Candidate Analysis"])

with tab1:
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("🎓 Candidate Evaluation Matrix")
        st.markdown("""
**Criteria** | **Points**
---|---
Master’s degree or higher | 10
Bachelor’s degree | 7
Diploma/Associate degree | 4
High school diploma or less | 2
10+ years experience | 10
5–10 years experience | 7
2–5 years experience | 5
<2 years experience | 3
High performance | 10
Above average performance | 7
Average performance | 5
Limited performance | 2
""")

    with col2:
        st.subheader("📈 Score-to-Step Matrix")
        st.markdown("""
| Score Range | Step Interval | Placement |
|-------------|----------------|-----------|
| 25–30 | Steps 12–15 | Top Range |
| 20–24 | Steps 9–11 | Mid-Upper Range |
| 15–19 | Steps 6–8 | Mid Range |
| 10–14 | Steps 3–5 | Lower-Mid Range |
| <10   | Steps 1–2 | Bottom Range |
""")

with tab2:
    st.subheader("Step 1: Candidate & Position Details")
    colA, colB = st.columns([1, 2])

    with colA:
        name = st.text_input("👤 Candidate Name")
        title = st.text_input("🏷️ Position Title (for equity comparison)")
        grade = st.text_input("🎖️ Position Grade (e.g., E-6)")

    st.subheader("Step 2: Upload Documents")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        uploaded_cv = st.file_uploader("📄 CV", type=["pdf", "docx"])
    with col2:
        uploaded_jd = st.file_uploader("📝 Job Description", type=["pdf", "docx"])
    with col3:
        uploaded_interviews = st.file_uploader("🗒️ Interview Sheets", type=["pdf", "docx"], accept_multiple_files=True)
    with col4:
        uploaded_equity = st.file_uploader("📊 Internal Equity Excel", type=["xlsx"])
    with col5:
        uploaded_scale = st.file_uploader("💼 Salary Scale (xlsx)", type=["xlsx"])

    # Load salary scale once and keep in session
    if uploaded_scale is not None:
        scale_df, scale_err = load_salary_scale(uploaded_scale)
        if scale_err:
            st.error(scale_err)
        else:
            st.session_state['salary_scale_df'] = scale_df
            st.success("Salary scale loaded.")
    elif 'salary_scale_df' not in st.session_state:
        st.session_state['salary_scale_df'] = None

    # Initialize stored AI scores (persist between reruns)
    if "ai_scores" not in st.session_state:
        st.session_state.ai_scores = {"educationScore": 0, "experienceScore": 0, "performanceScore": 0}

    st.subheader("Step 3: AI Evaluation + Manual Adjustment")
    run_ai = st.button("🔎 Run AI Analysis", help="Evaluates CV, JD, and Interview (if provided).")
    if run_ai:
        if not (uploaded_cv and uploaded_jd):
            st.warning("Please upload both the CV and Job Description to run AI analysis.")
        else:
            with st.spinner("🔍 Evaluating CV & JD..."):
                time.sleep(1)
                scores = mock_parse_cv_and_jd()

            if uploaded_interviews:
                with st.spinner("🧠 Evaluating Interview Sheet(s)..."):
                    time.sleep(1)
                    perf_scores = [mock_parse_interview_sheet()["performanceScore"] for _ in uploaded_interviews]
                    scores["performanceScore"] = max(perf_scores) if perf_scores else 0
            else:
                scores.setdefault("performanceScore", 0)

            st.session_state.ai_scores = scores
            st.success("AI analysis completed. You can manually adjust scores below.")

    education_score = st.slider("🎓 Education Score (Editable)", 0, 10, int(st.session_state.ai_scores.get("educationScore", 0)))
    experience_score = st.slider("💼 Experience Score (Editable)", 0, 10, int(st.session_state.ai_scores.get("experienceScore", 0)))
    performance_score = st.slider("🚀 Performance Score (Editable)", 0, 10, int(st.session_state.ai_scores.get("performanceScore", 0)))

    total_score = education_score + experience_score + performance_score
    interval_options, placement = get_step_interval(total_score)

    st.markdown(f"""
### 🎯 Candidate Scoring Summary
- Education: **{education_score}/10**
- Experience: **{experience_score}/10**
- Performance: **{performance_score}/10**
- Total Score: **{total_score}/30**
- Suggested Step Interval: **{interval_options} → {placement}**
""")

    selected_step = st.selectbox("✅ Select Final Step", interval_options)

    # Auto-fill AI-Recommended Salary from scale
    auto_rec = 0
    band_letter, grade_num = parse_band_grade(grade)
    scale_df = st.session_state.get('salary_scale_df', None)
    if scale_df is not None and band_letter and grade_num and selected_step:
        match = scale_df[
            (scale_df['band'] == band_letter) &
            (scale_df['grade'] == grade_num) &
            (scale_df['step'] == int(selected_step))
        ]
        if not match.empty:
            auto_rec = float(match.iloc[0]['amount'])

    # Keep user overrides stable across reruns
    if 'ai_rec_salary' not in st.session_state:
        st.session_state.ai_rec_salary = auto_rec
    if auto_rec and (st.session_state.ai_rec_salary == 0):
        st.session_state.ai_rec_salary = auto_rec

    st.subheader("Step 4: Salary Recommendation")
    budget_threshold = st.number_input("💰 Budget Threshold (AED)", step=500, value=0)
    recommended_salary = st.number_input("🤖 AI-Recommended Salary (AED)", step=500, value=float(st.session_state.ai_rec_salary or 0))
    st.session_state.ai_rec_salary = recommended_salary
    final_salary = st.number_input("✅ Final Recommended Salary (AED)", step=500, value=0)

    st.subheader("Step 5: Internal Equity Analysis")
    if uploaded_equity and title and final_salary == 0:
        st.info("ℹ️ Please enter the Final Recommended Salary (AED) above to perform the equity analysis.")

    df_peers = None
    if uploaded_equity and title and final_salary > 0:
        df_peers, error = load_filtered_equity_data(uploaded_equity, title)
        if error:
            st.error(error)
        elif df_peers is not None and df_peers.empty:
            st.warning(f"No matching peers found for position title: '{title}'")
        elif df_peers is not None:
            df_peers = pd.concat([
                df_peers,
                pd.DataFrame([{
                    "id": "Candidate",
                    "positionTitle": title,
                    "hireDate": "N/A",
                    "compRate": final_salary
                }])
            ], ignore_index=True)

            st.dataframe(df_peers)

            avg = df_peers["compRate"].mean()
            min_val = df_peers["compRate"].min()
            max_val = df_peers["compRate"].max()
            st.markdown(f"**Equity Range for '{title}':** Min AED {min_val:,.0f} | Avg AED {avg:,.0f} | Max AED {max_val:,.0f}")

            chart = alt.Chart(df_peers).mark_bar().encode(
                x=alt.X("id:N", title="Employee ID"),
                y=alt.Y("compRate:Q", title="Compensation (AED)"),
                color=alt.condition(alt.datum.id == "Candidate", alt.value("#FF8C00"), alt.value("#2a5c88")),
                tooltip=["id", "compRate"]
            ).properties(title="💼 Compensation Comparison", width=700, height=350)
            st.altair_chart(chart, use_container_width=True)

            excel_data = convert_df_to_excel(df_peers)
            st.download_button(
                label="📥 Download Equity Data (.xlsx)",
                data=excel_data,
                file_name="equity_analysis_filtered.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    st.subheader("Step 6: Budget Status")
    if final_salary and budget_threshold:
        if final_salary <= budget_threshold:
            st.success(f"✅ Within Budget (AED {final_salary:,.0f})")
        else:
            st.error(f"❌ Out of Budget (AED {final_salary:,.0f})")

    st.subheader("Step 7: Final HR Comments & Export")
    hr_comments = st.text_area("📝 HR Final Comments")

    if st.button("📤 Generate Final Summary"):
        summary = f"""
📌 Final Recommendation Summary

Candidate Name: {name}
Position Title: {title}
Grade: {grade}

Total Score: {total_score}/30 → Step Interval: {interval_options} → Placement: {placement}
Selected Step: {selected_step}
AI-Recommended Salary: AED {recommended_salary:,.0f}
Final Recommended Salary: AED {final_salary:,.0f}
Budget Threshold: AED {budget_threshold:,.0f}
Budget Status: {'Within' if final_salary <= budget_threshold else 'Out of'} Budget

HR Final Comments:
{hr_comments}
""".strip()

        st.text_area("📋 Final Summary", summary, height=250)
        st.download_button("📤 Download Final Summary (.txt)", data=summary, file_name="salary_summary.txt")

        doc = generate_word_report(
            name, title, grade,
            education_score, experience_score, performance_score,
            total_score, interval_options, placement, selected_step,
            recommended_salary, final_salary, budget_threshold,
            hr_comments, df_peers if isinstance(df_peers, pd.DataFrame) else None
        )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Download Full Report (.docx)",
            data=buffer,
            file_name=f"{(name or 'candidate').replace(' ', '_').lower()}_evaluation_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
